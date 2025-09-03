import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, logging
from docx import Document
from tqdm import tqdm
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import os, sys, time, gc

# ---------------------------
# CONFIGURATION
# ---------------------------
INPUT_DOCX = "./src/input.docx"
OUTPUT_MD = "./src/rewritten.md"
WARNINGS_LOG = "./src/warnings.txt"

CHUNK_SIZE = 80
MODEL_NAME = "TinyLlama/TinyLlama-1.1B-Chat-v1.0"
MODEL_DIR = "./model"
DEFAULT_FONT_SIZE = 12

# ---------------------------
# ENABLE LOGGING
# ---------------------------
logging.set_verbosity_error()
torch.set_printoptions(precision=3, sci_mode=False)

# ---------------------------
# LOAD / CACHE MODEL
# ---------------------------
if not os.path.exists(MODEL_DIR):
    print("=====================================================")
    print(f" Downloading model {MODEL_NAME} from Hugging Face...")
    os.makedirs(MODEL_DIR, exist_ok=True)

tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, cache_dir=MODEL_DIR)
model = AutoModelForCausalLM.from_pretrained(
    MODEL_NAME,
    cache_dir=MODEL_DIR,
    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
    low_cpu_mem_usage=True
)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

# ---------------------------
# HELPERS
# ---------------------------
def chunk_text(text, size=CHUNK_SIZE):
    """Split long text into smaller chunks."""
    words = text.split()
    return [" ".join(words[i:i+size]) for i in range(0, len(words), size)]

def llama_generate(prompt):
    """Generate text with low memory usage."""
    inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=256).to(device)

    with torch.no_grad():
        output_ids = model.generate(
            **inputs,
            max_new_tokens=200,
            do_sample=True,
            top_p=0.9,
            temperature=0.7
        )

    text = tokenizer.decode(output_ids[0], skip_special_tokens=True)

    # Free memory
    del inputs, output_ids
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()

    return text

def plagiarism_risk(original, rewritten):
    """Light plagiarism risk check using cosine similarity."""
    try:
        vec = TfidfVectorizer().fit([original, rewritten])
        sim = cosine_similarity(vec.transform([original]), vec.transform([rewritten]))[0][0]
        del vec
        gc.collect()
        return round(sim * 100, 2)
    except Exception:
        return None

def rewrite_text(text):
    if not text.strip():
        return text
    prompt = (
        f"Rewrite this paragraph in plagiarism-free English. "
        f"Use synonyms and natural wording, preserve meaning:\n\n{text}\n\nRewritten:"
    )
    return llama_generate(prompt).strip()

def get_font_size(para):
    """Get font size, default if missing."""
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    return DEFAULT_FONT_SIZE

def is_plain_paragraph(para):
    """Skip text inside shapes/textboxes."""
    parent_tag = para._element.getparent().tag.split("}")[-1]
    return parent_tag == "body"

def process_table(table, out_file):
    """Convert table to markdown."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if rows:
        header = rows[0]
        sep = "| " + " | ".join(["---"] * (len(rows[0].split('|')) - 2)) + " |"
        out_file.write("\n".join([header, sep] + rows[1:]) + "\n\n")

# ---------------------------
# MAIN PROCESSOR
# ---------------------------
def process_docx(input_file, output_file, warnings_file):
    doc = Document(input_file)
    prev_font_size = None

    with tqdm(total=len(doc.paragraphs) + len(doc.tables),
              desc="📖 Processing DOCX", unit="block") as pbar, \
         open(output_file, "w", encoding="utf-8") as f, \
         open(warnings_file, "w", encoding="utf-8") as warn:

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                pbar.update(1)
                continue

            if not is_plain_paragraph(para):
                warn.write(f"⚠️ Skipped text inside shape/box: \"{text[:60]}...\"\n")
                pbar.update(1)
                continue

            font_size = get_font_size(para)
            style_name = para.style.name.lower()

            if prev_font_size is not None and font_size != prev_font_size:
                f.write("\n")

            if "heading" in style_name:
                level = style_name.replace("heading", "").strip()
                level = int(level) if level.isdigit() else 2
                f.write("#" * level + " " + text + "\n\n")

            else:  # Paragraph
                if len(text.split()) > CHUNK_SIZE:
                    chunks = chunk_text(text)
                    rewritten_chunks = []
                    for chunk in chunks:
                        try:
                            rewritten_chunks.append(rewrite_text(chunk))
                        except Exception as e:
                            warn.write(f"⚠️ Rewrite failed: \"{chunk[:60]}...\" ({e})\n")
                    rewritten = " ".join(rewritten_chunks)
                else:
                    try:
                        rewritten = rewrite_text(text)
                    except Exception as e:
                        warn.write(f"⚠️ Rewrite failed: \"{text[:60]}...\" ({e})\n")
                        rewritten = text

                risk = plagiarism_risk(text, rewritten)
                if risk is not None:
                    warn.write(f"📊 Risk {risk}% | Original: \"{text[:60]}...\"\n")

                f.write(rewritten + "\n\n")

            prev_font_size = font_size
            pbar.update(1)

        for table in doc.tables:
            process_table(table, f)
            pbar.update(1)

    print(f"\n✅ Saved: {output_file}")
    print(f"📑 Warnings report: {warnings_file}")

# ---------------------------
# RUN
# ---------------------------
if __name__ == "__main__":
    process_docx(INPUT_DOCX, OUTPUT_MD, WARNINGS_LOG)