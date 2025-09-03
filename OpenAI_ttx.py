import os
import openai
from docx import Document
from tqdm import tqdm

# ---------------------------
# CONFIGURATION
# ---------------------------
INPUT_DOCX = "./src/input.docx"
OUTPUT_MD = "./src/rewritten.md"
CHUNK_SIZE = 500                   # Words per chunk for API
OPENAI_MODEL = "gpt-3.5-turbo"     # or "gpt-4"
openai.api_key = os.getenv("OPENAI_API_KEY")

# ---------------------------
# HELPERS
# ---------------------------

def chunk_text(text, size=CHUNK_SIZE):
    """Split text into word chunks for API"""
    words = text.split()
    return [" ".join(words[i:i+size]) for i in range(0, len(words), size)]

def detect_heading_level(para):
    """Detect heading based on style name or font size"""
    style_name = getattr(para.style, "name", "Normal").lower()
    if "heading 1" in style_name:
        return 1
    elif "heading 2" in style_name:
        return 2
    elif "heading 3" in style_name:
        return 3
    # fallback: check max font size
    sizes = [run.font.size.pt for run in para.runs if run.font.size]
    if sizes:
        max_size = max(sizes)
        if max_size >= 20:
            return 1
        elif max_size >= 16:
            return 2
        elif max_size >= 14:
            return 3
    return 0  # normal paragraph

def wrap_font_size(para, text):
    """Wrap text in <span> with font size for HTML/Markdown"""
    sizes = [run.font.size.pt for run in para.runs if run.font.size]
    avg_size = int(sum(sizes)/len(sizes)) if sizes else 12
    return f'<span style="font-size:{avg_size}pt">{text}</span>'

def rewrite_chunk(text, is_heading=False, is_list=False):
    """Send text to OpenAI and return rewritten text"""
    if is_heading:
        prompt = f"Rephrase this heading in simple, human-like, plagiarism-free words:\n\n{text}\n\nNew Heading:"
    elif is_list:
        prompt = f"Rephrase this bullet point in easy words with synonyms:\n\n{text}\n\nNew Bullet:"
    else:
        prompt = f"Rephrase this paragraph in simple natural English, plagiarism-free:\n\n{text}\n\nRewritten:"

    response = openai.ChatCompletion.create(
        model=OPENAI_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=1024
    )
    rewritten = response['choices'][0]['message']['content'].strip()
    return rewritten

def process_paragraph(para):
    """Process paragraph with headings, lists, font size preservation"""
    text = para.text.strip()
    if not text:
        return ""

    # Detect heading level
    level = detect_heading_level(para)
    is_heading = level > 0
    is_list = "list" in getattr(para.style, "name", "").lower()

    # Split on soft line breaks
    lines = text.split("\n")
    rewritten_lines = []

    # Progress bar for rewriting each line
    for line in tqdm(lines, desc="✏️ Rewriting paragraph", leave=False):
        if line.strip():
            rewritten = rewrite_chunk(line, is_heading=is_heading, is_list=is_list)
            if is_heading:
                rewritten_lines.append(wrap_font_size(para, rewritten))
            else:
                rewritten_lines.append(rewritten)

    if is_heading:
        return "#"*level + " " + " ".join(rewritten_lines) + "\n"
    elif is_list:
        return "- " + " ".join(rewritten_lines) + "\n"
    else:
        return " ".join(rewritten_lines) + "\n"

def process_table(table):
    """Convert Word tables to Markdown tables"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if rows:
        header = rows[0]
        separator = "| " + " | ".join(["---"]*(len(rows[0].split("|"))-2)) + " |"
        return "\n".join([header, separator]+rows[1:]) + "\n"
    return ""

# ---------------------------
# MAIN DOCX PROCESSOR
# ---------------------------
def process_docx(input_file, output_file):
    doc = Document(input_file)
    md_lines = []
    elements = list(doc.element.body)
    total = len(elements)
    table_idx = 0

    with tqdm(total=total, desc="📖 Processing DOCX", unit="block") as pbar:
        for block in elements:
            tag = block.tag.split("}")[-1]
            if tag == "p":
                para = doc.paragraphs[elements.index(block)]
                if para.text.strip():
                    if para.text.startswith("$") and para.text.endswith("$"):
                        md_lines.append(para.text+"\n")
                    else:
                        rewritten = process_paragraph(para)
                        md_lines.append(rewritten)
            elif tag == "tbl":
                table_obj = doc.tables[table_idx]
                print("📊 Processing table...")
                md_lines.append(process_table(table_obj))
                table_idx += 1
            elif tag in ["drawing", "pict"]:
                print("🖼️ Found image placeholder")
                md_lines.append("![Image](image-placeholder.png)\n")
            else:
                md_lines.append("<!-- Other element preserved -->\n")
            pbar.update(1)
            pbar.set_postfix({"Remaining": total - pbar.n})

    with open(output_file, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    print(f"\n✅ Rewritten Markdown saved as {output_file}")

# ---------------------------
# RUN
# ---------------------------
if __name__=="__main__":
    process_docx(INPUT_DOCX, OUTPUT_MD)