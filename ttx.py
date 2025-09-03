import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, logging
from docx import Document
from tqdm import tqdm
import os
import sys
import time

# ---------------------------
# CONFIGURATION
# ---------------------------
INPUT_DOCX = "./src/input.docx"         # Input Word file
OUTPUT_MD = "./src/rewritten.md"        # Output Markdown file
WARNINGS_LOG = "./src/warnings.txt"     # Warnings report
CHUNK_SIZE = 120                        # Words per chunk
MODEL_NAME = "TinyLlama/TinyLlama-1.1B-Chat-v1.0"
MODEL_DIR = "./model"                   # Local save directory for model
DEFAULT_FONT_SIZE = 12                  # Default font size (pt)

# ---------------------------
# ENABLE LOGGING
# ---------------------------
logging.set_verbosity_info()
logging.enable_explicit_format()

os.environ["TORCH_SHOW_CPP_STACKTRACES"] = "1"
torch.set_printoptions(precision=4, sci_mode=False)

# ---------------------------
# LOAD / CACHE MODEL
# ---------------------------
if not os.path.exists(MODEL_DIR):
    print("=====================================================")
    print(" Downloading model from Hugging Face...")

tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, cache_dir=MODEL_DIR)
model = AutoModelForCausalLM.from_pretrained(
    MODEL_NAME,
    cache_dir=MODEL_DIR,
    dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
    device_map="auto" if torch.cuda.is_available() else None
)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

# ---------------------------
# HELPERS
# ---------------------------
def chunk_text(text, size=CHUNK_SIZE):
    """Split long text into smaller chunks."""
    words = text.split()
    return [" ".join(words[i:i + size]) for i in range(0, len(words), size)]

def llama_generate(prompt, stream=True):
    """Generate rewritten text from TinyLlama with streaming logs."""
    print("\n🧠 Model generating...")
    inputs = tokenizer(prompt, return_tensors="pt").to(device)

    if stream:
        generated = inputs["input_ids"]
        prev_text = ""

        for _ in range(512):  # max tokens
            with torch.no_grad():
                outputs = model(input_ids=generated)
                next_token_logits = outputs.logits[:, -1, :]
                next_token_id = torch.argmax(next_token_logits, dim=-1).unsqueeze(0)

            generated = torch.cat((generated, next_token_id), dim=1)

            decoded_so_far = tokenizer.decode(generated[0], skip_special_tokens=True)
            new_text = decoded_so_far[len(prev_text):]

            if new_text:
                sys.stdout.write(new_text)
                sys.stdout.flush()
                time.sleep(0.01)

            prev_text = decoded_so_far

            if next_token_id.item() == tokenizer.eos_token_id:
                break

        print("\n✅ Streaming complete.\n")
        return prev_text

    else:
        outputs = model.generate(
            **inputs,
            max_length=512,
            do_sample=True,
            top_p=0.9,
            temperature=0.7
        )
        return tokenizer.decode(outputs[0], skip_special_tokens=True)

def rewrite_text(text):
    """Rewrites only paragraph text with logging."""
    if not text.strip():
        return text

    prompt = f"Rephrase this paragraph in natural easy English, plagiarism-free:\n\n{text}\n\nRewritten:"
    print(f"\n🔄 Rewriting: {text[:80]}...")
    return llama_generate(prompt, stream=True).strip()

def get_font_size(para):
    """Extract font size of first run in a paragraph (default if missing)."""
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    return DEFAULT_FONT_SIZE

def is_plain_paragraph(para):
    """Check if paragraph is plain (not inside shape/textbox)."""
    parent_tag = para._element.getparent().tag.split("}")[-1]
    return parent_tag == "body"

def process_table(table):
    """Convert Word tables into Markdown tables."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if rows:
        header = rows[0]
        separator = "| " + " | ".join(["---"] * (len(rows[0].split('|')) - 2)) + " |"
        return "\n".join([header, separator] + rows[1:]) + "\n\n"
    return ""

# ---------------------------
# MAIN PROCESSOR
# ---------------------------
def process_docx(input_file, output_file, warnings_file):
    doc = Document(input_file)
    prev_font_size = None

    with tqdm(total=len(doc.paragraphs) + len(doc.tables),
              desc="📖 Processing DOCX", unit="block") as pbar, \
         open(output_file, "w", encoding="utf-8") as f, \
         open(warnings_file, "w", encoding="utf-8") as warn:

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                pbar.update(1)
                continue

            if not is_plain_paragraph(para):
                warn.write(f"⚠️ Skipped text inside shape/box: \"{text[:60]}...\"\n")
                pbar.update(1)
                continue

            font_size = get_font_size(para)
            style_name = para.style.name.lower()

            # Insert new line if font size changes or new paragraph
            if prev_font_size is not None and font_size != prev_font_size:
                f.write("\n")

            if "heading" in style_name:
                level = style_name.replace("heading", "").strip()
                level = int(level) if level.isdigit() else 2
                f.write("#" * level + " " + text + "\n\n")

            else:  # Normal paragraph → rewrite
                if len(text.split()) > CHUNK_SIZE:
                    chunks = chunk_text(text)
                    rewritten_chunks = []
                    for chunk in chunks:
                        try:
                            rewritten_chunks.append(rewrite_text(chunk))
                        except Exception as e:
                            warn.write(f"⚠️ Rewrite failed for chunk: \"{chunk[:60]}...\" ({e})\n")
                    f.write(" ".join(rewritten_chunks) + "\n\n")
                else:
                    try:
                        rewritten = rewrite_text(text)
                        f.write(rewritten + "\n\n")
                    except Exception as e:
                        warn.write(f"⚠️ Rewrite failed for paragraph: \"{text[:60]}...\" ({e})\n")

            prev_font_size = font_size
            pbar.update(1)

        # Process tables
        for table in doc.tables:
            f.write(process_table(table))
            pbar.update(1)

    print(f"\n✅ Rewritten Markdown saved as {output_file}")
    print(f"📑 Warnings saved as {warnings_file}")

# ---------------------------
# RUN
# ---------------------------
if __name__ == "__main__":
    process_docx(INPUT_DOCX, OUTPUT_MD, WARNINGS_LOG)